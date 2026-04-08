import os
import uuid
import json
import logging
from fastapi import FastAPI, UploadFile, File, WebSocket, WebSocketDisconnect, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from backend.config import get_settings
from backend.orchestrator import ShieldWallOrchestrator
from backend.models.schemas import ShieldWallJobState
from backend.policy_store.indexer import index_policies
from backend.middleware.logging_middleware import StructuredLoggingMiddleware
from backend.health import router as health_router

logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)

app = FastAPI()
app.include_router(health_router)
app.add_middleware(StructuredLoggingMiddleware)

settings = get_settings()

allowed_origins = [
    "http://localhost:5173",
    "http://localhost:5174",
    "http://localhost:5175",
    os.getenv("FRONTEND_ORIGIN", ""),
]
allowed_origins = [o for o in allowed_origins if o]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

orchestrator = ShieldWallOrchestrator(settings)

@app.on_event("startup")
async def startup():
    os.makedirs("/tmp/shield_wall", exist_ok=True)
    orchestrator.initialize()

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    ext = file.filename.split('.')[-1].lower()
    if ext not in ['xlsx', 'csv', 'pdf', 'docx', 'txt']:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    job_id = str(uuid.uuid4())
    job_dir = f"/tmp/shield_wall/{job_id}"
    os.makedirs(job_dir, exist_ok=True)
    file_path = f"{job_dir}/questionnaire.{ext}"
    
    content = await file.read()
    if len(content) > settings.max_file_size_mb * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large")
        
    with open(file_path, "wb") as f:
        f.write(content)
        
    orchestrator.jobs[job_id] = ShieldWallJobState(job_id=job_id, status="pending")
    return {"job_id": job_id}

@app.websocket("/ws/{job_id}")
async def websocket_endpoint(websocket: WebSocket, job_id: str):
    await websocket.accept()
    
    if job_id not in orchestrator.jobs:
        await websocket.send_text(json.dumps({"error": "Job not found"}))
        await websocket.close()
        return
        
    # Get the file path
    job_dir = f"/tmp/shield_wall/{job_id}"
    files = os.listdir(job_dir)
    if not files:
        await websocket.close()
        return
    file_path = os.path.join(job_dir, files[0])
    
    async def ws_callback(event):
        try:
            await websocket.send_text(event.model_dump_json())
        except Exception:
            logger.warning(f"WebSocket send failed for job {job_id}")

    try:
        await orchestrator.run_pipeline(job_id, file_path, ws_callback)
    except WebSocketDisconnect:
        logger.info(f"Client disconnected for job {job_id}")
    except Exception as e:
        logger.error(f"Pipeline error for job {job_id}: {e}")

@app.get("/api/result/{job_id}")
async def get_result(job_id: str):
    if job_id not in orchestrator.jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    return orchestrator.jobs[job_id].model_dump()

@app.get("/api/costs/{job_id}")
async def get_costs(job_id: str):
    if job_id not in orchestrator.jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = orchestrator.jobs[job_id]
    entries = getattr(job, "cost_entries", [])
    total_cost = sum(e.estimated_cost_usd for e in entries)
    return {
        "job_id": job_id,
        "entries": [e.model_dump() for e in entries],
        "total_cost_usd": total_cost
    }

@app.get("/api/export/{job_id}")
async def export_job(job_id: str):
    if job_id not in orchestrator.jobs:
        raise HTTPException(status_code=404, detail="Job not found")

    job = orchestrator.jobs[job_id]
    if job.status != "complete":
        raise HTTPException(status_code=400, detail="Job not complete")

    export_path = f"/tmp/shield_wall/{job_id}/completed_questionnaire.docx"

    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()
    doc.add_heading("Shield-Wall — Completed Security Questionnaire", level=0)

    summary = doc.add_paragraph()
    summary.add_run(f"Total questions: {job.result.total_questions}  |  ")
    summary.add_run(f"High confidence: {job.result.high_confidence}  |  ")
    summary.add_run(f"Drift alerts: {job.result.drift_alerts}  |  ")
    summary.add_run(f"Needs review: {job.result.needs_review}")
    summary.paragraph_format.space_after = Pt(12)

    for ans in job.result.answers:
        q_text = ans.answer_text
        # Find original question text from questionnaire
        orig = ""
        if job.questionnaire:
            for sq in job.questionnaire.questions:
                if sq.id == ans.question_id:
                    orig = sq.original_text
                    break

        heading = doc.add_heading(f"Q{ans.question_id}", level=2)
        if orig:
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(orig)
            q_run.italic = True
            q_run.font.color.rgb = RGBColor(100, 100, 100)

        a_para = doc.add_paragraph()
        a_para.add_run(ans.answer_text)

        conf_para = doc.add_paragraph()
        conf_run = conf_para.add_run(f"Confidence: {ans.confidence.upper()}")
        conf_run.bold = True
        if ans.drift_detected:
            drift_para = doc.add_paragraph()
            drift_run = drift_para.add_run(f"DRIFT DETECTED: {ans.drift_detail}")
            drift_run.font.color.rgb = RGBColor(220, 50, 50)

    doc.save(export_path)
    return FileResponse(
        export_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="completed_questionnaire.docx"
    )

@app.post("/api/policies/reindex")
async def reindex_policies():
    orchestrator.initialize()
    return {"status": "success"}
